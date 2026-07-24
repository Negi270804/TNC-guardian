import os
import sys
import time
import logging
from typing import TypedDict, Optional
import numpy as np
from PIL import Image, ImageOps, ImageEnhance, ImageFilter

logger = logging.getLogger("app.services.ocr_service")

class ExtractionResult(TypedDict):
    text: str
    page_count: int
    word_count: int

class OCRService:
    _reader = None

    @classmethod
    def get_reader(cls):
        """Lazy-loaded, cached EasyOCR reader instance to conserve RAM/VRAM resource pools."""
        if cls._reader is None:
            try:
                import easyocr
                import torch
                from app import config
                use_gpu = torch.cuda.is_available() if config.OCR_USE_GPU else False
                langs = [lang.strip() for lang in config.OCR_LANGUAGES.split(",") if lang.strip()]
                logger.info(f"[OCR SERVICE] Initializing EasyOCR Reader with languages {langs} (GPU Enabled: {use_gpu})...")
                cls._reader = easyocr.Reader(langs, gpu=use_gpu)
            except Exception as e:
                logger.error(f"[OCR SERVICE] Failed to initialize EasyOCR library: {str(e)}", exc_info=True)
                raise RuntimeError(f"OCR Reader engine failed to start: {str(e)}")
        return cls._reader

    @classmethod
    async def extract_text(cls, file_path: str, file_type: str) -> ExtractionResult:
        """Extracts text from files according to type parameters, falling back to OCR when needed."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Source file not found at path: {file_path}")

        file_ext = file_type.lower().strip('.')
        
        # 1. TXT Documents
        if file_ext == "txt":
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
                return {
                    "text": text,
                    "page_count": 1,
                    "word_count": len(text.split())
                }
            except Exception as e:
                raise RuntimeError(f"Failed to read TXT file: {str(e)}")

        # 2. DOCX Documents
        elif file_ext == "docx":
            try:
                import docx
                doc = docx.Document(file_path)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                # Parse tables too for completeness
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text for cell in row.cells]
                        full_text.append(" | ".join(row_text))
                
                text = "\n".join(full_text)
                return {
                    "text": text,
                    "page_count": 1,
                    "word_count": len(text.split())
                }
            except Exception as e:
                raise RuntimeError(f"Failed to extract DOCX file text: {str(e)}")

        # 3. PDF Documents
        elif file_ext == "pdf":
            try:
                import pdfplumber
                text = ""
                page_count = 0
                
                with pdfplumber.open(file_path) as pdf:
                    page_count = len(pdf.pages)
                    for i, page in enumerate(pdf.pages):
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text += page_text + "\n"
                        else:
                            # PDF Page has no selectable text, fall back to OCR on the page image
                            try:
                                pil_img = page.to_image(resolution=150).original.convert('RGB')
                                img_arr = np.array(pil_img)
                                reader = cls.get_reader()
                                ocr_results = reader.readtext(img_arr, detail=0)
                                ocr_text = " ".join(ocr_results)
                                if ocr_text.strip():
                                    text += ocr_text + "\n"
                            except Exception as ocr_err:
                                logger.warning(f"[OCR WARNING] Failed to OCR PDF Page {i+1}: {str(ocr_err)}")
                                text += f"[Page {i+1} OCR Extraction Failure]\n"
                
                return {
                    "text": text,
                    "page_count": page_count,
                    "word_count": len(text.split())
                }
            except Exception as e:
                raise RuntimeError(f"Failed to parse PDF file layout: {str(e)}")

        # 4. Image Documents (PNG, JPG, JPEG, WEBP, BMP)
        elif file_ext in ["png", "jpg", "jpeg", "webp", "bmp"]:
            start_time = time.time()
            try:
                # Validate file exists and get size
                file_size = os.path.getsize(file_path)
                logger.info(f"Image received: {file_path}")
                logger.info(f"Image size: {file_size} bytes")

                # Validate file size
                if file_size <= 0:
                    raise ValueError("Invalid empty file uploaded.")
                
                from app import config
                pro_limit_bytes = config.PRO_PLAN_UPLOAD_LIMIT_MB * 1024 * 1024
                if file_size > pro_limit_bytes:
                    raise ValueError(f"File size exceeds the maximum limit of {config.PRO_PLAN_UPLOAD_LIMIT_MB} MB.")

                # Open and validate image format
                try:
                    with Image.open(file_path) as pil_img:
                        pil_img.load()
                        img_format = pil_img.format.upper() if pil_img.format else ""
                        width, height = pil_img.size
                except Exception as img_err:
                    raise ValueError(f"Invalid or corrupted image file: {str(img_err)}")

                logger.info(f"Image validated. Format: {img_format}, Dimensions: {width}x{height}")
                
                # Check supported format
                if img_format not in ["PNG", "JPEG", "JPG", "MPO", "WEBP", "BMP"]:
                    raise ValueError(f"Unsupported image format in file: {img_format}")

                with Image.open(file_path) as pil_img:
                    # Upscale image if it is small or has small dimensions to improve OCR accuracy
                    width, height = pil_img.size
                    if width < 1200 or height < 1200:
                        scale = 2
                        if width < 500 or height < 500:
                            scale = 3
                        pil_img = pil_img.resize((width * scale, height * scale), Image.Resampling.LANCZOS)

                    # Auto orientation if required (exif_transpose)
                    pil_img = ImageOps.exif_transpose(pil_img)

                    # Handle transparency (RGBA, LA, or P with transparency)
                    if pil_img.mode in ("RGBA", "LA") or (pil_img.mode == "P" and "transparency" in pil_img.info):
                        # Composite on a solid white background
                        bg = Image.new("RGBA", pil_img.size, (255, 255, 255, 255))
                        composite = Image.alpha_composite(bg, pil_img.convert("RGBA"))
                        processed_rgb = composite.convert("RGB")
                    else:
                        processed_rgb = pil_img.convert("RGB")

                    # Convert to grayscale
                    gray_img = processed_rgb.convert("L")

                    # Increase contrast using autocontrast
                    gray_img = ImageOps.autocontrast(gray_img)
                    # Enhance contrast further
                    enhancer = ImageEnhance.Contrast(gray_img)
                    gray_img = enhancer.enhance(2.0)

                    # Noise reduction (Median Filter size 3)
                    gray_img = gray_img.filter(ImageFilter.MedianFilter(size=3))

                    # Otsu Threshold calculation for soft thresholding band
                    img_arr = np.array(gray_img)
                    try:
                        pixel_counts = np.bincount(img_arr.ravel(), minlength=256)
                        total_pixels = img_arr.size
                        p = pixel_counts / float(total_pixels)
                        q = np.cumsum(p)
                        bins = np.arange(256)
                        mu = np.cumsum(bins * p)
                        mu_g = mu[-1]
                        with np.errstate(divide='ignore', invalid='ignore'):
                            denom = q * (1.0 - q)
                            variance = (mu_g * q - mu) ** 2 / denom
                        variance[np.isnan(variance)] = 0.0
                        thresh = np.argmax(variance)
                    except Exception:
                        thresh = 127

                    # Adaptive Soft Thresholding (retains gradients for anti-aliasing edges)
                    low = max(0, thresh - 90)
                    high = min(255, thresh + 90)
                    span = high - low if high > low else 1
                    binary_arr = np.clip((img_arr - low) * 255.0 / span, 0, 255).astype(np.uint8)

                    # Dynamic Inversion (ensure black text on white background)
                    mean_val = np.mean(binary_arr)
                    if mean_val < 127:
                        binary_arr = 255 - binary_arr

                    preprocessed_img = Image.fromarray(binary_arr)

                    # Sharpen edges using Unsharp Mask
                    preprocessed_img = preprocessed_img.filter(ImageFilter.UnsharpMask(radius=2, percent=150, threshold=3))

                    # OCR Execution
                    logger.info(f"OCR started: format={file_ext}")
                    reader = cls.get_reader()
                    
                    ocr_results = reader.readtext(np.array(preprocessed_img), detail=1)

                # Preserve Line Breaks
                blocks = []
                for res in ocr_results:
                    if len(res) == 3:
                        bbox, text_val, conf = res
                    elif len(res) == 2:
                        bbox, text_val = res
                        conf = 1.0
                    else:
                        continue
                    if not text_val or not text_val.strip():
                        continue
                        
                    xs = [pt[0] for pt in bbox]
                    ys = [pt[1] for pt in bbox]
                    ymin, ymax = min(ys), max(ys)
                    xmin, xmax = min(xs), max(xs)
                    
                    blocks.append({
                        "ymin": ymin,
                        "ymax": ymax,
                        "xmin": xmin,
                        "xmax": xmax,
                        "y_center": (ymin + ymax) / 2.0,
                        "height": ymax - ymin,
                        "text": text_val
                    })
                
                # Group into lines based on vertical overlap
                lines = []
                for block in sorted(blocks, key=lambda b: b["y_center"]):
                    placed = False
                    for line in lines:
                        line_ymin = min(b["ymin"] for b in line)
                        line_ymax = max(b["ymax"] for b in line)
                        line_height = line_ymax - line_ymin
                        
                        overlap_ymin = max(block["ymin"], line_ymin)
                        overlap_ymax = min(block["ymax"], line_ymax)
                        overlap = max(0.0, overlap_ymax - overlap_ymin)
                        
                        min_h = min(block["height"], line_height)
                        if min_h > 0 and (overlap / min_h) > 0.4:
                            line.append(block)
                            placed = True
                            break
                            
                    if not placed:
                        lines.append([block])
                
                # Sort horizontally within each line
                line_texts = []
                for line in lines:
                    sorted_line = sorted(line, key=lambda b: b["xmin"])
                    line_text = " ".join(b["text"] for b in sorted_line)
                    line_texts.append(line_text)
                    
                text = "\n".join(line_texts).strip()

                # If empty, return standard fallback
                if not text:
                    text = "No readable text detected in the uploaded image."

                elapsed_time = time.time() - start_time
                logger.info(f"OCR completed: {file_path}")
                logger.info(f"Characters extracted: {len(text)}")
                logger.info(f"Processing time: {elapsed_time:.2f} seconds")

                return {
                    "text": text,
                    "page_count": 1,
                    "word_count": len(text.split())
                }
            except Exception as e:
                elapsed_time = time.time() - start_time
                logger.error(f"OCR execution failed for image {file_path} after {elapsed_time:.2f}s: {str(e)}", exc_info=True)
                raise RuntimeError(f"Failed to OCR image file: {str(e)}")

        else:
            raise ValueError(f"Unsupported file type extension passed for text extraction: '{file_ext}'")
